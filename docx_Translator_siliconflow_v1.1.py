import docx
from openai import OpenAI
import os
import re

# Recommended: Securely get API key from environment variable, use hardcoded backup if not set
# Powershell command: [Environment]::SetEnvironmentVariable("siliconflow_API_KEY", "sk-zzzzzz", "User")
api_key = os.environ.get("siliconflow_API_KEY")
if not api_key:
    api_key = "sk-xxxxxx"

docx_path = "1.docx"
output_path = "1-CN.docx"
font_modified = "Microsoft YaHei Light"

# Initialize OpenAI client
client = OpenAI(api_key=api_key, base_url="https://api.siliconflow.cn/v1")


def split_into_sentences(text):
    pattern = r"(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s"
    sentences = re.split(pattern, text)
    return [s.strip() for s in sentences if s.strip()]


def translate_text(text, target_language="zh-CN"):
    if not text or len(text.strip()) < 2:
        return text

    sentences = split_into_sentences(text)
    translated_sentences = []

    for sentence in sentences:
        try:
            response = client.chat.completions.create(
                model="Qwen/Qwen2-7B-Instruct",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            f"You are a professional, authentic machine translation engine. "
                            f"Translate the following sentence to {target_language}. "
                            "Output translation directly without any additional text. "
                            "Don't translate IT terms. "
                            "Don't translate the words beginning with 'Forti'. "
                            "Don't translate words in single quotes ('output', 'spoke'). "
                            "Keep the original words unchanged which you can't recognize."
                        ),
                    },
                    {"role": "user", "content": sentence},
                ],
                temperature=0.2,
            )
            translated_sentences.append(response.choices[0].message.content.strip())
        except Exception as e:
            print(f"Translation error for sentence '{sentence[:50]}...': {e}")
            translated_sentences.append(sentence)

    return " ".join(translated_sentences)


def safe_set_font(run):
    try:
        run.font.name = font_modified
    except AttributeError:
        pass


def translate_paragraph(paragraph, target_language):
    try:
        full_text = paragraph.text
        if not full_text.strip():
            return

        translated_text = translate_text(full_text, target_language)

        # Clear existing runs
        for _ in range(len(paragraph.runs)):
            p = paragraph._element
            p.remove(p.r_lst[0])

        # Add a new run with the translated text
        new_run = paragraph.add_run(translated_text)
        safe_set_font(new_run)

        print(f"Original Text: {full_text[:50]}...")
        print(f"Translated Text: {translated_text[:50]}...")
    except Exception as e:
        print(f"Error translating paragraph: {e}")


def translate_table(table, target_language):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translate_paragraph(paragraph, target_language)


def translate_docx(docx_path, target_language="zh-CN", output_path="translated.docx"):
    try:
        doc = docx.Document(docx_path)

        # Translate main document content
        for paragraph in doc.paragraphs:
            translate_paragraph(paragraph, target_language)

        # Translate tables
        for table in doc.tables:
            translate_table(table, target_language)

        # Translate headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                translate_paragraph(header, target_language)
            for footer in section.footer.paragraphs:
                translate_paragraph(footer, target_language)

        # Translate document properties
        if doc.core_properties.title:
            doc.core_properties.title = translate_text(
                doc.core_properties.title, target_language
            )
        if doc.core_properties.subject:
            doc.core_properties.subject = translate_text(
                doc.core_properties.subject, target_language
            )

        doc.save(output_path)
        print(f"Translated DOCX saved to {output_path}")
    except Exception as e:
        print(f"Error processing document: {e}")


# Execute translation
translate_docx(docx_path=docx_path, target_language="zh-CN", output_path=output_path)
